import re
import os
import fitz  # PyMuPDF
import json
import docx
import torch
from transformers import BertTokenizer, BertForSequenceClassification,BertForTokenClassification,BertTokenizerFast

CATEGORY_KEYWORDS = {
    "基本信息": ["个人信息", "个人资料", "基本信息", "个人主页", "个人简介", "自我评价", "个人优势", "主题描述"],
    "教育背景": ["教育背景", "教育经历", "学历信息", "学术背景", "学习经历"],
    "工作/项目经历": ["工作经历", "工作经验", "职业经历", "实习经历",  "主要经历", "校园经历", "最近工作", "项目经验", "项目经历", "项目成果", "实践项目", "科研项目", "科研经历","科研竞赛"],
    "个人能力": ["技能", "个人技能", "特长", "个人特长", "奖项荣誉", "获奖经历", "个人荣誉","证书", "证书情况", "技能证书", "语言能力", "计算机技能", "专业技能"],
    "其他信息": ["兴趣爱好"]
}
# 构建正则表达式模式
patterns = []
for category, keywords in CATEGORY_KEYWORDS.items():
    patterns.append(r'(?:\n|^|[\u3002\s])({})(?:[：:、\s]|$)'.format('|'.join(keywords)))

pattern = re.compile('|'.join(patterns), flags=re.UNICODE)

def split_resume(text):
    """将简历文本分割成不同的部分"""
    # 文本预处理
    text = re.sub(r'[\r\n]+', '\n', text)  # 统一换行符
    text = re.sub(r'[ 　]+', ' ', text)     # 合并连续空格
    text = text.strip()
    
    # 处理教育背景关键词
    education_keywords = '|'.join(CATEGORY_KEYWORDS['教育背景'])
    text = re.sub(r'([^\n])([^：:、\s])({})'.format(education_keywords), r'\1\n\2\3', text)
    
    # 处理其他关键词
    for category, keywords in CATEGORY_KEYWORDS.items():
        if category != '教育背景':  
            keyword_pattern = '|'.join(keywords)
            text = re.sub(r'([^\n])([^：:、\s])({})'.format(keyword_pattern), r'\1\n\2\3', text)
    
    # 处理没有换行符的关键词
    for category, keywords in CATEGORY_KEYWORDS.items():
        for keyword in keywords:
            keyword_pattern = r'(?<!\n)({})'.format(keyword)
            text = re.sub(keyword_pattern, r'\n\1', text)
    
    last_end = 0
    last_category = None
    blocks = []
    header_content = ""

    # 使用全局的pattern变量
    for match in pattern.finditer(text):
        start = match.start()
        if start > last_end:
            # 如果是第一个块且没有类别，则视为header
            if last_category is None and not blocks:
                header_content = text[last_end:start].strip()
                # 如果header内容长度大于10，则将其作为基本信息的一部分
                if len(header_content) > 10:
                    blocks.append({
                        "category": "基本信息",
                        "content": header_content,
                        "start_pos": last_end,
                        "end_pos": start
                    })
                else:
                    blocks.append({
                        "category": None,
                        "content": header_content,
                        "start_pos": last_end,
                        "end_pos": start
                    })
            else:
                blocks.append({
                    "category": last_category,
                    "content": text[last_end:start].strip(),
                    "start_pos": last_end,
                    "end_pos": start
                })

        # 获取匹配到的关键词类别
        matched_keyword = match.group().strip() 
        category = next((cat for cat, keys in CATEGORY_KEYWORDS.items()
                        if any(k in matched_keyword for k in keys)), None)

        last_category = category
        last_end = match.end()

    # 处理最后一段
    if last_end < len(text):
        blocks.append({
            "category": last_category,
            "content": text[last_end:].strip(),
            "start_pos": last_end,
            "end_pos": len(text)
        })

    # 合并相同类别的板块
    merged_blocks = []
    category_blocks = {}
    
    for block in blocks:
        category = block["category"]
        if category not in category_blocks:
            category_blocks[category] = block.copy()
        else:
            # 合并内容，避免重复
            existing_content = category_blocks[category]["content"]
            new_content = block["content"]
            if new_content not in existing_content:
                category_blocks[category]["content"] = existing_content + "\n" + new_content
            category_blocks[category]["end_pos"] = block["end_pos"]
    
    # 将分类后的块添加到结果中
    merged_blocks = list(category_blocks.values())
    
    # 后处理：确保每个块都有内容
    merged_blocks = [block for block in merged_blocks if block["content"].strip()]
    
    # 处理header内容：如果header内容长度大于10，将其合并到基本信息板块
    if any(block["category"] == "基本信息" for block in merged_blocks):
        basic_info_block = next(block for block in merged_blocks if block["category"] == "基本信息")
        header_block = next((block for block in merged_blocks if block["category"] is None and len(block["content"]) > 10), None)
        
        if header_block:
            # 合并header内容到基本信息板块，避免重复
            if header_block["content"] not in basic_info_block["content"]:
                basic_info_block["content"] = header_block["content"] + "\n" + basic_info_block["content"]
                basic_info_block["start_pos"] = header_block["start_pos"]
            # 移除已合并的header_block
            merged_blocks = [block for block in merged_blocks if block != header_block]
    
    # 确保所有分类都存在，为空分类添加'无'标记
    all_categories = ["基本信息", "教育背景", "工作/项目经历", "个人能力", "其他信息"]
    existing_categories = {block["category"] for block in merged_blocks if block["category"]}
    
    for category in all_categories:
        if category not in existing_categories:
            merged_blocks.append({
                "category": category,
                "content": "无",
                "start_pos": -1,
                "end_pos": -1
            })

    return merged_blocks

def extract_name_from_sections(sections):
    """从简历分段结果中提取姓名"""
    # 初始化NER模型
    model_path = 'models/ner/ner_model'
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    tokenizer = BertTokenizerFast.from_pretrained(model_path)
    model = BertForTokenClassification.from_pretrained(model_path).to(device)
    
    # 从基本信息中提取姓名
    if '基本信息' in sections:
        profile_text = sections['基本信息']
        # 使用NER模型预测
        model.eval()
        tokens = tokenizer.tokenize(profile_text)
        encoding = tokenizer(
            tokens,
            is_split_into_words=True,
            max_length=128,
            padding='max_length',
            truncation=True,
            return_tensors='pt'
        ).to(device)
        
        with torch.no_grad():
            outputs = model(encoding['input_ids'], attention_mask=encoding['attention_mask'])
            preds = torch.argmax(outputs.logits, dim=-1)
        
        # 对齐预测结果
        word_ids = encoding.word_ids()
        predictions = []
        previous_word_idx = None
        for i, (word_idx, pred) in enumerate(zip(word_ids, preds[0])):
            if word_idx is not None and word_idx != previous_word_idx:
                predictions.append((tokens[word_idx], model.config.id2label[pred.item()]))
                previous_word_idx = word_idx
        
        # 提取人名实体
        name_entities = []
        current_entity = []
        current_type = None
        
        for token, label in predictions:
            if label.startswith('B-') and label[2:] == 'PER':
                if current_entity:
                    name_entities.append(''.join(current_entity))
                    current_entity = []
                current_entity.append(token)
                current_type = 'PER'
            elif label.startswith('I-') and label[2:] == 'PER' and current_type == 'PER':
                current_entity.append(token)
            else:
                if current_entity:
                    name_entities.append(''.join(current_entity))
                    current_entity = []
                current_type = None
        
        if current_entity:
            name_entities.append(''.join(current_entity))
        
        # 如果找到人名实体，返回第一个
        if name_entities:
            return name_entities[0]
    
    return "未知"

def extract_pdf(pdfpath):
    """提取PDF文本内容"""
    text = ""
    try:
        doc = fitz.open(pdfpath)
        for page in doc:
            text += page.get_text()
        
        doc.close()
        text = text.strip()
        return text
    except Exception as e:
        print(f"PDF提取错误: {str(e)}")
        return ""
    
def extract_docx(filepath):
    """提取DOCX文件内容"""
    try:
        doc = docx.Document(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"DOCX提取错误: {str(e)}")
        return ""

def is_valid_file(filepath):
    """检查文件是否有效"""
    try:
        if filepath.lower().endswith('.pdf'):
            # 验证PDF文件
            doc = fitz.open(filepath)
            if len(doc) > 0:
                doc.close()
                return True
            doc.close()
            return False
        elif filepath.lower().endswith('.docx'):
            # 验证DOCX文件
            doc = docx.Document(filepath)
            # 如果能成功打开文档，说明文件有效
            return True
        return False
    except Exception as e:
        print(f"文件验证错误: {str(e)}")
        return False

def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx'}

def split_sentences(text):
    """将文本分割成句子"""
    text = text.replace('⚫', '').split(" \n")
    return [line.replace('\n', '') for line in text if line]

def classify_with_bert(text):
    """使用BERT模型对文本进行分类"""
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    print(f'使用设备: {device}')

    # 使用相对路径
    model_dir = "models/classify/best_model"  

    print(f'从 {model_dir} 加载模型...')

    try:
        # 加载模型和分词器
        tokenizer = BertTokenizer.from_pretrained(model_dir, local_files_only=True)
        model = BertForSequenceClassification.from_pretrained(model_dir, local_files_only=True)
        model.to(device)
    except Exception as e:
        print(f"加载模型失败: {e}")
        return None

    sentences = split_sentences(text)
    # 类别名称
    class_names = ['基本信息', '教育背景', '工作/项目经历', '个人能力', '其他信息']
    sections = {category: "" for category in class_names}
    
    print(f"分割出的句子数量: {len(sentences)}")
    
    # 对每个句子进行预测
    print("\n开始对句子进行分类...")
    for i, sentence in enumerate(sentences, 1):
        if not sentence.strip():
            continue
            
        try:
            # 编码与预测
            model.eval()
            encoding = tokenizer(
                sentence,
                max_length=512,
                padding='max_length',
                truncation=True,
                return_tensors='pt'
            ).to(device)
            
            with torch.no_grad():
                outputs = model(**encoding)
                logits = outputs.logits if hasattr(outputs, 'logits') else outputs[0]
                
            probabilities = torch.nn.functional.softmax(logits, dim=1)
            predicted_class = torch.argmax(probabilities, dim=1).item()
            
            # 拼接结果
            separator = '\n' if len(sentence) > 20 else ';'
            sections[class_names[predicted_class]] += f"{sentence}{separator}"
            
        except Exception as e:
            print(f"处理句子时出错: {str(e)}")
            sections['其他信息'] += f"\n{sentence}"
    
    print("\n分类完成!")
    return sections

def bert_parser(text, filename=None):
    """使用BERT模型解析简历文本"""
    try:
        # 使用BERT进行分类
        print("\n开始BERT分类...")
        bert_sections = classify_with_bert(text)
        if bert_sections is None:
            return None, "BERT分类失败"

        # 将BERT分类结果转换为统一的blocks格式
        blocks = []
        for category in ['基本信息', '教育背景', '工作/项目经历', '个人能力', '其他信息']:
            content = bert_sections.get(category, '').strip()
            blocks.append({
                "category": category,
                "content": content if content else "无"
            })

        # 从基本信息中提取姓名
        name = "未知"
        if '基本信息' in bert_sections:
            # 构造正确的数据格式
            sections_dict = {'基本信息': bert_sections['基本信息']}
            name = extract_name_from_sections(sections_dict)

        result = {
            'filename': filename,
            'name': name,
            'blocks': blocks
        }
        
        return result, None
        
    except Exception as e:
        print(f"\nBERT解析过程中发生错误: {str(e)}")
        return None, str(e)

def rule_parser(text, filename=None):
    """使用规则解析简历文本"""
    try:
        # 解析简历结构
        print("\n开始规则解析...")
        sections_blocks = split_resume(text)
        print(f"解析出的段落数量: {len(sections_blocks)}")
        
        # 将sections转换为字典格式
        sections_dict = {}
        for block in sections_blocks:
            if block['category']:
                sections_dict[block['category']] = block['content']

        # 从基本信息中提取姓名
        name = "未知"
        if '基本信息' in sections_dict:
            name = extract_name_from_sections(sections_dict)

        # 创建统一的blocks格式
        blocks = []
        for category in ['基本信息', '教育背景', '工作/项目经历', '个人能力', '其他信息']:
            content = sections_dict.get(category, '').strip()
            blocks.append({
                "category": category,
                "content": content if content else "无"
            })

        result = {
            'filename': filename,
            'name': name,
            'blocks': blocks
        }
        
        return result, None
        
    except Exception as e:
        print(f"\n规则解析过程中发生错误: {str(e)}")
        return None, str(e)

def save_resume(data, output_dir="output"):
    """保存结构化数据到JSON文件"""
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{os.path.splitext(data['filename'])[0]}.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return output_path

def show_resume(data):
    """显示简历解析结果"""
    print("\n=== 简历内容 ===")
    print(f"文件: {data['filename']}")
    print(f"时间: {data['timestamp']}\n")

    print("=== 原始文本 ===")
    print("-" * 60)
    print(data['text'][:500] + "..." if len(data['text']) > 500 else data['text'])
    print("-" * 60)

    print("\n=== 结构化数据 ===")
    for block in data['blocks']:
        title = block['category'] or "未分类"
        print(f"\n[{title}]")
        print("-" * 40)
        print(block['content'])
        print("-" * 40)

    # 显示BERT分类结果
    if 'bert_sections' in data:
        print("\n=== BERT分类结果 ===")
        for category, content in data['bert_sections'].items():
            if content.strip():
                print(f"\n[{category}]")
                print("-" * 40)
                print(content.strip())
                print("-" * 40)
